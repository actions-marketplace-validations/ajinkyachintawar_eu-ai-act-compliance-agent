from __future__ import annotations

import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

from backend.models.schemas import (
    ComplianceReport,
    ObligationStatus,
    RiskTier,
)

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------

BLUE_DARK = RGBColor(0x1F, 0x49, 0x7D)   # EU-ish dark blue — headings
BLUE_MID  = RGBColor(0x2E, 0x75, 0xB6)   # mid blue — sub-headings
RED_SOFT  = RGBColor(0xC0, 0x00, 0x00)   # [INFORMATION REQUIRED] text
GREEN     = RGBColor(0x37, 0x86, 0x44)   # MET
RED       = RGBColor(0xC0, 0x00, 0x00)   # NOT MET
AMBER     = RGBColor(0xBF, 0x82, 0x00)   # UNCLEAR
GREY_LIGHT = "D9D9D9"                     # table header fill (hex string for XML)
GREEN_FILL = "E2EFDA"
RED_FILL   = "FCE4D6"
AMBER_FILL = "FFF2CC"


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background colour via raw XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc: Document, color: str = "2E75B6") -> None:
    """Add a thin coloured horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def _add_page_break(doc: Document) -> None:
    doc.add_paragraph().runs  # paragraph
    doc.paragraphs[-1].add_run().add_break()


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _heading1(doc: Document, text: str, numbering: str = "") -> None:
    label = f"{numbering}  {text}" if numbering else text
    p = doc.add_heading(label, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE_DARK
        run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def _heading2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLUE_MID
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _body(doc: Document, text: str) -> None:
    """Add body paragraph, highlighting [INFORMATION REQUIRED] in red."""
    pattern = re.compile(r"(\[INFORMATION REQUIRED[^\]]*\])")
    parts = pattern.split(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(11)
        if i % 2 == 1:  # odd indices are the matched [INFORMATION REQUIRED] spans
            run.font.color.rgb = RED_SOFT
            run.bold = True


def _label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    lbl = p.add_run(f"{label}: ")
    lbl.bold = True
    lbl.font.size = Pt(11)
    val = p.add_run(value)
    val.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _build_cover(doc: Document, report: ComplianceReport) -> None:
    sy = report.system
    cl = report.classification

    # Large title
    doc.add_paragraph()  # top padding
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EU AI ACT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE_DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Annex IV — Technical Documentation")
    run2.font.size = Pt(18)
    run2.font.color.rgb = BLUE_MID

    _add_horizontal_rule(doc)

    doc.add_paragraph()

    # System name box
    sys_p = doc.add_paragraph()
    sys_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sys_p.add_run(sy.name)
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()

    # Classification badge
    tier_colors = {
        RiskTier.HIGH_RISK: "C00000",
        RiskTier.PROHIBITED: "7B0000",
        RiskTier.LIMITED_RISK: "BF8200",
        RiskTier.MINIMAL_RISK: "378644",
    }
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = badge.add_run(f"Risk Classification: {cl.tier.value}  ({cl.confidence:.0%} confidence)")
    br.bold = True
    br.font.size = Pt(13)
    br.font.color.rgb = RGBColor.from_string(tier_colors.get(cl.tier, "1F497D"))

    doc.add_paragraph()
    _add_horizontal_rule(doc)

    # Meta info
    meta = [
        ("Sector", sy.sector.value),
        ("Decision type", sy.decision_type.value),
        ("Generated", report.generated_at.strftime("%d %B %Y, %H:%M UTC")),
        ("Document status", "DRAFT — for internal review only"),
    ]
    for label, value in meta:
        _label_value(doc, label, value)

    doc.add_paragraph()
    disc = doc.add_paragraph(report.disclaimer)
    disc.paragraph_format.space_after = Pt(0)
    for run in disc.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_system_summary(doc: Document, report: ComplianceReport) -> None:
    _heading1(doc, "General Description of the AI System", "1.")
    _add_horizontal_rule(doc)
    av = report.annex_iv_draft
    sy = report.system
    _body(doc, av.system_description)

    _heading2(doc, "System Details")
    fields = [
        ("Name", sy.name),
        ("Use case", sy.use_case),
        ("Sector", sy.sector.value),
        ("Inputs", sy.inputs),
        ("Outputs", sy.outputs),
        ("Affected persons", sy.affected_persons),
        ("Decision type", sy.decision_type.value),
        ("Replaces existing practice", sy.existing_practices or "New capability"),
    ]
    for label, value in fields:
        _label_value(doc, label, value)


def _build_intended_purpose(doc: Document, report: ComplianceReport) -> None:
    _heading1(doc, "Intended Purpose and Deployment Context", "2.")
    _add_horizontal_rule(doc)
    _body(doc, report.annex_iv_draft.intended_purpose)


def _build_performance(doc: Document, report: ComplianceReport) -> None:
    _heading1(doc, "Performance Metrics", "3.")
    _add_horizontal_rule(doc)
    _body(doc, report.annex_iv_draft.performance_metrics)


def _build_risk_management(doc: Document, report: ComplianceReport) -> None:
    _heading1(doc, "Risk Management Summary (Article 9)", "4.")
    _add_horizontal_rule(doc)
    _body(doc, report.annex_iv_draft.risk_management_summary)

    _heading2(doc, "Article 6 Exception Analysis")
    ex = report.article6_exception
    p = doc.add_paragraph()
    r = p.add_run("Qualifies for exception: ")
    r.bold = True
    r.font.size = Pt(11)
    val = p.add_run("YES" if ex.qualifies else "NO")
    val.bold = True
    val.font.size = Pt(11)
    val.font.color.rgb = GREEN if ex.qualifies else RED
    _body(doc, ex.reasoning)


def _build_data_governance(doc: Document, report: ComplianceReport) -> None:
    _heading1(doc, "Data Governance Notes (Article 10)", "5.")
    _add_horizontal_rule(doc)
    _body(doc, report.annex_iv_draft.data_governance_notes)


def _build_obligations_table(doc: Document, report: ComplianceReport) -> None:
    _heading1(doc, "Obligations Checklist", "6.")
    _add_horizontal_rule(doc)

    ob = report.obligations
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_after = Pt(8)
    for label, count, color in [
        ("✓ MET", ob.total_met, GREEN),
        ("✗ NOT MET", ob.total_not_met, RED),
        ("? UNCLEAR", ob.total_unclear, AMBER),
    ]:
        r = summary_p.add_run(f"  {label}: {count}  ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = color

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Article", "Obligation", "Status", "Evidence"]):
        _set_cell_bg(cell, GREY_LIGHT)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BLUE_DARK

    status_fills = {
        ObligationStatus.MET: GREEN_FILL,
        ObligationStatus.NOT_MET: RED_FILL,
        ObligationStatus.UNCLEAR: AMBER_FILL,
    }
    status_colors = {
        ObligationStatus.MET: GREEN,
        ObligationStatus.NOT_MET: RED,
        ObligationStatus.UNCLEAR: AMBER,
    }

    for o in ob.obligations:
        row = table.add_row().cells
        fill = status_fills[o.status]

        row[0].paragraphs[0].add_run(o.article).font.size = Pt(9)
        row[1].paragraphs[0].add_run(o.title).font.size = Pt(9)
        _set_cell_bg(row[2], fill)
        status_run = row[2].paragraphs[0].add_run(o.status.value)
        status_run.bold = True
        status_run.font.size = Pt(9)
        status_run.font.color.rgb = status_colors[o.status]
        row[3].paragraphs[0].add_run(o.evidence or "—").font.size = Pt(9)

    # Column widths (approximate — docx doesn't guarantee them)
    col_widths = [Inches(1.0), Inches(2.2), Inches(1.0), Inches(2.8)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]


# ---------------------------------------------------------------------------
# Header / footer
# ---------------------------------------------------------------------------

def _set_header_footer(doc: Document, system_name: str) -> None:
    section = doc.sections[0]

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(f"EU AI Act — Annex IV | {system_name}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    r.italic = True

    # Footer with page numbers
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Page ").font.size = Pt(9)
    # Page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    fp.add_run(" | DRAFT — not legal advice").font.size = Pt(9)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_annex_iv_docx(report: ComplianceReport) -> bytes:
    """
    Generate an Annex IV Technical Documentation .docx from a ComplianceReport.
    Returns the document as bytes (ready to stream from a FastAPI endpoint).
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Cover page
    _build_cover(doc, report)
    _add_page_break(doc)

    # Sections
    _build_system_summary(doc, report)
    _build_intended_purpose(doc, report)
    _build_performance(doc, report)
    _build_risk_management(doc, report)
    _build_data_governance(doc, report)
    _build_obligations_table(doc, report)

    # Header / footer on all pages
    _set_header_footer(doc, report.system.name)

    # Serialise to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
